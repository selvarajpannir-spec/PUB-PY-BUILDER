"""
Docs+ Document Analyser
-----------------------
Local preview + top tags per file. No AI.

Java calls:
  analyse_one(path) -> JSON object string for one file
  list is done on Java side; Python only analyses.

Rules:
  - One file at a time (caller loops)
  - Skip-by-name is done in the UI/table layer
  - Scan cap 50 is enforced by the UI loop
"""

import os
import re
import json
from collections import Counter

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

MAX_PDF_PAGES = 25
MIN_WORD_LEN = 4
MIN_COUNT = 3
TOP_TAGS = 5
PREVIEW_CHARS = 160

STOP_WORDS = {
    "the", "and", "to", "a", "of", "in", "for", "is", "on", "that", "by",
    "this", "with", "i", "you", "it", "not", "or", "be", "are", "from",
    "at", "as", "your", "all", "have", "an", "my", "has", "we", "he",
    "they", "she", "but", "so", "if", "will", "do", "can", "their", "which",
    "about", "there", "would", "who", "get", "when", "make", "like", "time",
    "just", "him", "know", "take", "people", "into", "year", "good",
    "some", "could", "them", "see", "other", "than", "then", "now", "look",
    "only", "come", "its", "over", "think", "also", "back", "use", "two",
    "how", "our", "work", "first", "well", "way", "even", "new", "want",
    "because", "any", "these", "give", "day", "most", "us", "was", "were",
    "been", "being", "had", "did", "am", "his", "her", "me", "no", "yes",
    "may", "shall", "should", "must", "might", "please", "using", "used",
    "made", "such", "each", "more", "very", "much", "many", "own", "same",
    "too", "page", "pages", "click", "http", "https", "www", "pdf", "doc",
    "docx", "file", "files", "com", "org", "index", "contents", "chapter",
    "section", "figure", "table", "appendix", "total", "date", "name",
    "type", "details", "description", "others", "out", "upon", "within",
    "without", "under", "above", "below", "between", "through", "during",
    "before", "after", "again", "further", "once", "here", "where", "why",
    "both", "few", "nor", "don", "ltd", "pte", "inc", "etc", "ie", "eg",
}


def extract_text(path):
    low = path.lower()
    try:
        if low.endswith(".pdf"):
            if not HAS_PDF:
                return "", "pypdf not installed"
            reader = PdfReader(path)
            parts = []
            for page in reader.pages[:MAX_PDF_PAGES]:
                t = page.extract_text()
                if t:
                    parts.append(t)
            return " ".join(parts), None

        if low.endswith((".txt", ".md")):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), None

        if low.endswith(".xlsx"):
            if not HAS_XLSX:
                return "", "openpyxl not installed"
            wb = load_workbook(path, data_only=True, read_only=True)
            parts = []
            for ws in wb.worksheets[:3]:
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i > 200:
                        break
                    for cell in row:
                        if cell is not None:
                            parts.append(str(cell))
            wb.close()
            return " ".join(parts), None

        if low.endswith(".docx"):
            if not HAS_DOCX:
                return "", "python-docx not installed"
            doc = DocxDocument(path)
            parts = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    parts.append(para.text)
            for table in doc.tables[:20]:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip() if cell.text else ""
                        if t:
                            parts.append(t)
            return " ".join(parts), None
    except Exception as e:
        return "", str(e)
    return "", "unsupported type"


def filename_tokens(filename):
    base = os.path.splitext(filename)[0].lower()
    return set(re.findall(r"[a-z]{3,}", base))


def tokenize(text):
    return re.findall(r"\b[a-zA-Z]{3,}\b", text.lower())


def top_tags(text, filename):
    words = tokenize(text)
    filtered = [w for w in words if w not in STOP_WORDS and len(w) >= MIN_WORD_LEN]
    if not filtered:
        return []
    counts = Counter(filtered)
    candidates = [(w, c) for w, c in counts.items() if c >= MIN_COUNT]
    if not candidates:
        candidates = counts.most_common(20)
    fname = filename_tokens(filename)

    def score(item):
        w, c = item
        return (c + (4 if w in fname else 0), len(w))

    ranked = sorted(candidates, key=score, reverse=True)
    tags = []
    for w, _ in ranked:
        if w not in tags:
            tags.append(w)
        if len(tags) >= TOP_TAGS:
            break
    return tags


def preview_text(text):
    if not text:
        return "(no text extracted)"
    t = re.sub(r"\s+", " ", text).strip()
    parts = re.split(r"(?<=[.!?])\s+", t)
    good = []
    for s in parts:
        s = s.strip()
        if len(s) < 40:
            continue
        if re.match(r"^(page\s*\d+|contents|chapter\s*\d+)", s, re.I):
            continue
        letters = sum(ch.isalpha() for ch in s)
        if letters < len(s) * 0.5:
            continue
        good.append(s)
        if len(good) >= 2:
            break
    out = " ".join(good) if good else t[:PREVIEW_CHARS]
    out = out.strip()
    if len(out) > PREVIEW_CHARS:
        out = out[: PREVIEW_CHARS - 1].rsplit(" ", 1)[0] + "…"
    return out


def analyse_one(path):
    """
    Analyse a single file. Return JSON string:
    { ok, name, path, preview, tags, note, error? }
    """
    try:
        if not path or not os.path.isfile(path):
            return json.dumps({
                "ok": False,
                "name": os.path.basename(path or ""),
                "path": path or "",
                "preview": "",
                "tags": [],
                "note": "File not found",
                "error": "File not found",
            }, ensure_ascii=False)

        filename = os.path.basename(path)
        text, err = extract_text(path)
        if err and not text:
            return json.dumps({
                "ok": False,
                "name": filename,
                "path": path,
                "preview": "",
                "tags": [],
                "note": "Could not read: " + err,
                "error": err,
            }, ensure_ascii=False)

        if not text or len(text.strip()) < 30:
            return json.dumps({
                "ok": False,
                "name": filename,
                "path": path,
                "preview": "(too little text)",
                "tags": [],
                "note": "(too little text)",
                "error": "too little text",
            }, ensure_ascii=False)

        prev = preview_text(text)
        tags = top_tags(text, filename)
        note = (prev + "  |  Tags: " + " · ".join(tags)) if tags else prev
        return json.dumps({
            "ok": True,
            "name": filename,
            "path": path,
            "preview": prev,
            "tags": tags,
            "note": note,
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({
            "ok": False,
            "name": os.path.basename(path or ""),
            "path": path or "",
            "preview": "",
            "tags": [],
            "note": str(e),
            "error": str(e),
        }, ensure_ascii=False)
