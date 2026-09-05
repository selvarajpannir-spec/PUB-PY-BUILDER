"""
PDF Generator — Android (Chaquopy).
- Text: {{COLUMN_NAME}}
- Optional photos from Excel embedded images (by row).
- {{PHOTO}} / {{IMAGE}} in Word table cell = box; image forced to fit cell size.
- Max 50 rows.
"""

import os
import re
import zipfile
import tempfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Twips
from docxcompose.composer import Composer
from openpyxl import load_workbook

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

IMAGE_PLACEHOLDERS = {"PHOTO", "IMAGE", "PIC", "PICTURE"}

# Used only when cell size cannot be read from the template
DEFAULT_WIDTH_CM = 3.5
DEFAULT_HEIGHT_CM = 4.0
MAX_ROWS = 50

NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
}


def paragraph_has_drawing(paragraph):
    for run in paragraph.runs:
        if run._element.findall(qn("w:drawing")):
            return True
    return False


def _clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _twips_to_cm(twips):
    try:
        return float(twips) * 2.54 / 1440.0
    except Exception:
        return None


def _emu_to_cm(val):
    try:
        if val is None:
            return None
        if hasattr(val, "cm"):
            return float(val.cm)
        return float(val) / 360000.0
    except Exception:
        return None


def _cell_box_cm(cell):
    """
    Read photo box size from the Word table cell.
    Prefer explicit tcW / trHeight (most reliable).
    """
    width_cm = None
    height_cm = None

    # 1) Cell width from tcPr/tcW (twips or pct — we use dxa/twips)
    try:
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                w_val = tcW.get(qn("w:w"))
                w_type = tcW.get(qn("w:type")) or "dxa"
                if w_val and w_type in ("dxa", "nil", None):
                    width_cm = _twips_to_cm(w_val)
    except Exception:
        pass

    # 2) python-docx cell.width
    if width_cm is None or width_cm < 0.4:
        try:
            if cell.width is not None:
                w = _emu_to_cm(cell.width)
                if w and w >= 0.4:
                    width_cm = w
        except Exception:
            pass

    # 3) Column grid width from table
    if width_cm is None or width_cm < 0.4:
        try:
            table = cell._tc.getparent().getparent()  # tr -> tbl
            # find cell index
            tr = cell._tc.getparent()
            tcs = list(tr)
            idx = tcs.index(cell._tc)
            grid = table.find(qn("w:tblGrid"))
            if grid is not None:
                cols = list(grid.findall(qn("w:gridCol")))
                if idx < len(cols):
                    g = cols[idx].get(qn("w:w"))
                    if g:
                        width_cm = _twips_to_cm(g)
        except Exception:
            pass

    # Row height from trHeight
    try:
        tr = cell._tc.getparent()
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None:
            trHeight = trPr.find(qn("w:trHeight"))
            if trHeight is not None:
                val = trHeight.get(qn("w:val"))
                if val:
                    h = _twips_to_cm(val)
                    if h and h >= 0.4:
                        height_cm = h
    except Exception:
        pass

    if width_cm is None or width_cm < 0.4:
        width_cm = DEFAULT_WIDTH_CM
    if height_cm is None or height_cm < 0.4:
        height_cm = DEFAULT_HEIGHT_CM

    # Small margin so image doesn't clip cell borders (~6%)
    width_cm *= 0.94
    height_cm *= 0.94

    # Clamp to sane range for mobile docs
    width_cm = max(1.0, min(width_cm, 10.0))
    height_cm = max(1.0, min(height_cm, 12.0))

    return width_cm, height_cm


def _fit_image_file(src_path, max_width_cm, max_height_cm, dpi=96):
    """
    Resize to fit inside box (keep aspect ratio).
    Returns (path, display_width_cm, display_height_cm).
    """
    if not HAS_PIL:
        return src_path, max_width_cm, max_height_cm
    try:
        max_w = max(1, int(max_width_cm / 2.54 * dpi))
        max_h = max(1, int(max_height_cm / 2.54 * dpi))
        img = PILImage.open(src_path)
        if img.mode not in ("RGB", "RGBA"):
            img = img.convert("RGBA" if "A" in img.getbands() else "RGB")

        img.thumbnail((max_w, max_h), PILImage.Resampling.LANCZOS)
        new_w, new_h = img.size
        out_w_cm = new_w / float(dpi) * 2.54
        out_h_cm = new_h / float(dpi) * 2.54

        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        img.save(tmp.name, format="PNG", optimize=True)
        tmp.close()
        return tmp.name, out_w_cm, out_h_cm
    except Exception:
        return src_path, max_width_cm, max_height_cm


def insert_image_in_paragraph(paragraph, image_path, width_cm=None, height_cm=None):
    """
    Always size the image to the Word template table cell (box).
    Excel image dimensions are ignored — only the template cell matters.
    The picture is forced to the cell width x height so it fills the box.
    """
    if width_cm is None or width_cm <= 0:
        width_cm = DEFAULT_WIDTH_CM
    if height_cm is None or height_cm <= 0:
        height_cm = DEFAULT_HEIGHT_CM

    _clear_paragraph_runs(paragraph)

    # Pre-resize pixels to roughly the target box (saves memory); display size
    # is always forced to the template cell below.
    fitted, _, _ = _fit_image_file(image_path, width_cm, height_cm)

    run = paragraph.add_run()
    try:
        # Force exact box size from the template table cell
        run.add_picture(fitted, width=Cm(width_cm), height=Cm(height_cm))
    except Exception:
        try:
            run.add_picture(fitted, width=Cm(width_cm))
        except Exception:
            try:
                run.add_picture(fitted)
            except Exception:
                paragraph.add_run("")

    if fitted != image_path:
        try:
            os.remove(fitted)
        except Exception:
            pass


def _is_image_token(full_text, mapping, image_keys):
    stripped = full_text.strip()
    for key in image_keys:
        token = "{{" + key + "}}"
        if stripped == token or token in full_text:
            return True
    for key in mapping:
        if key.strip().upper() in image_keys:
            token = "{{" + key + "}}"
            if stripped == token or token in full_text:
                return True
    return False


def replace_text_in_paragraph(
    paragraph, mapping, image_path=None, image_keys=None, box_cm=None
):
    if paragraph_has_drawing(paragraph):
        return
    runs = paragraph.runs
    full_text = "".join(run.text for run in runs) if runs else ""
    image_keys = image_keys or IMAGE_PLACEHOLDERS
    width_cm, height_cm = box_cm if box_cm else (DEFAULT_WIDTH_CM, DEFAULT_HEIGHT_CM)

    if image_path and _is_image_token(full_text, mapping, image_keys):
        insert_image_in_paragraph(paragraph, image_path, width_cm, height_cm)
        return

    if not full_text:
        return

    new_text = full_text
    for key, val in mapping.items():
        token = "{{" + key + "}}"
        if key.strip().upper() in image_keys:
            if image_path is None:
                new_text = new_text.replace(token, "")
            continue
        new_text = new_text.replace(token, str(val) if val is not None else "")

    if image_path is None:
        for key in image_keys:
            new_text = new_text.replace("{{" + key + "}}", "")

    if new_text == full_text:
        return

    font = runs[0].font if runs else None
    _clear_paragraph_runs(paragraph)
    new_run = paragraph.add_run(new_text)
    if font is not None:
        new_run.font.name = font.name
        new_run.font.size = font.size
        new_run.font.bold = font.bold
        new_run.font.italic = font.italic
        new_run.font.underline = font.underline
        try:
            if font.color and font.color.rgb:
                new_run.font.color.rgb = font.color.rgb
        except Exception:
            pass


def replace_in_table(table, mapping, image_path=None, image_keys=None):
    for row in table.rows:
        for cell in row.cells:
            box = _cell_box_cm(cell)
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(
                    paragraph, mapping,
                    image_path=image_path,
                    image_keys=image_keys,
                    box_cm=box,
                )


def replace_in_section(section, mapping, image_path=None, image_keys=None):
    try:
        for para in section.header.paragraphs:
            replace_text_in_paragraph(para, mapping, image_path, image_keys)
        for para in section.footer.paragraphs:
            replace_text_in_paragraph(para, mapping, image_path, image_keys)
        for table in section.header.tables:
            replace_in_table(table, mapping, image_path, image_keys)
        for table in section.footer.tables:
            replace_in_table(table, mapping, image_path, image_keys)
    except Exception:
        pass


def fill_document(doc, mapping, image_path=None, image_keys=None):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(
            paragraph, mapping, image_path=image_path, image_keys=image_keys
        )
    for table in doc.tables:
        replace_in_table(table, mapping, image_path=image_path, image_keys=image_keys)
    for section in doc.sections:
        replace_in_section(section, mapping, image_path=image_path, image_keys=image_keys)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(
        parse_xml(
            r'<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
        )
    )


def _anchor_row_openpyxl(image):
    try:
        anchor = image.anchor
        if hasattr(anchor, "_from") and anchor._from is not None:
            return int(anchor._from.row) + 1
        if hasattr(anchor, "from_") and anchor.from_ is not None:
            return int(anchor.from_.row) + 1
    except Exception:
        pass
    return None


def _extract_via_openpyxl(excel_path):
    result = {}
    try:
        wb = load_workbook(excel_path)
        ws = wb.active
        for img in getattr(ws, "_images", None) or []:
            row = _anchor_row_openpyxl(img)
            if row is None or row < 2:
                continue
            try:
                data = img._data()
            except Exception:
                continue
            if not data or not isinstance(data, (bytes, bytearray)):
                continue
            suffix = ".png"
            try:
                fmt = (getattr(img, "format", None) or "").lower()
                if fmt in ("jpeg", "jpg"):
                    suffix = ".jpg"
            except Exception:
                pass
            tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
            tmp.write(data)
            tmp.close()
            if row not in result:
                result[row] = tmp.name
        wb.close()
    except Exception:
        pass
    return result


def _extract_via_zip(excel_path):
    result = {}
    try:
        with zipfile.ZipFile(excel_path, "r") as zf:
            names = zf.namelist()
            drawing_files = [
                n for n in names
                if n.startswith("xl/drawings/drawing") and n.endswith(".xml")
            ]
            for drawing_path in drawing_files:
                rels_path = drawing_path.replace(
                    "xl/drawings/", "xl/drawings/_rels/"
                ) + ".rels"
                rid_to_media = {}
                if rels_path in names:
                    try:
                        rel_root = ET.fromstring(zf.read(rels_path))
                        for rel in rel_root:
                            rid = rel.get("Id")
                            target = rel.get("Target")
                            if rid and target and "media" in target:
                                media_name = target.replace("../", "xl/")
                                if not media_name.startswith("xl/"):
                                    media_name = os.path.normpath(
                                        "xl/drawings/" + target
                                    ).replace("\\", "/")
                                rid_to_media[rid] = media_name
                    except Exception:
                        pass

                try:
                    root = ET.fromstring(zf.read(drawing_path))
                except Exception:
                    continue

                for anchor in list(root):
                    tag = anchor.tag
                    if not (tag.endswith("twoCellAnchor") or tag.endswith("oneCellAnchor")):
                        continue
                    row = None
                    for child in anchor:
                        if child.tag.endswith("from"):
                            for sub in child:
                                if sub.tag.endswith("row") and sub.text is not None:
                                    try:
                                        row = int(sub.text) + 1
                                    except Exception:
                                        pass
                    if row is None or row < 2:
                        continue

                    embed = None
                    for el in anchor.iter():
                        if el.tag.endswith("blip"):
                            embed = el.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed is None:
                                embed = el.get("embed")
                            break
                    if not embed or embed not in rid_to_media:
                        continue

                    media_path = rid_to_media[embed]
                    if media_path not in names:
                        base = os.path.basename(media_path)
                        candidates = [n for n in names if n.endswith(base)]
                        if not candidates:
                            continue
                        media_path = candidates[0]

                    try:
                        data = zf.read(media_path)
                    except Exception:
                        continue
                    if not data:
                        continue

                    suffix = ".png"
                    low = media_path.lower()
                    if low.endswith(".jpg") or low.endswith(".jpeg"):
                        suffix = ".jpg"
                    elif low.endswith(".gif"):
                        suffix = ".gif"

                    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
                    tmp.write(data)
                    tmp.close()
                    if row not in result:
                        result[row] = tmp.name
    except Exception:
        pass
    return result


def extract_images_by_row(excel_path):
    result = _extract_via_openpyxl(excel_path)
    if result:
        return result
    return _extract_via_zip(excel_path)


def generate(excel_path, template_path, output_path, progress_callback=None):
    image_temps = []
    try:
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active
        headers = [str(cell.value).strip() for cell in ws[1] if cell.value is not None]

        rows = []
        row_numbers = []
        for excel_row_idx, row in enumerate(
            ws.iter_rows(min_row=2, values_only=True), start=2
        ):
            if all(cell is None or str(cell).strip() == "" for cell in row):
                break
            rows.append(row)
            row_numbers.append(excel_row_idx)
        wb.close()

        if not rows:
            return False, "No data rows found in Excel", 0

        truncated = False
        total_in_file = len(rows)
        if len(rows) > MAX_ROWS:
            rows = rows[:MAX_ROWS]
            row_numbers = row_numbers[:MAX_ROWS]
            truncated = True

        images_by_row = extract_images_by_row(excel_path)
        image_temps = list(images_by_row.values())

        total = len(rows)
        temp_files = []

        for i, row in enumerate(rows):
            mapping = {}
            for idx, h in enumerate(headers):
                mapping[h] = row[idx] if idx < len(row) else ""

            excel_row = row_numbers[i]
            image_path = images_by_row.get(excel_row)

            doc = Document(template_path)
            fill_document(doc, mapping, image_path=image_path)

            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            doc.save(tmp.name)
            temp_files.append(tmp.name)
            tmp.close()

            if progress_callback:
                try:
                    progress_callback(i + 1, total)
                except Exception:
                    pass

        master = Document(temp_files[0])
        composer = Composer(master)
        for tmp_path in temp_files[1:]:
            add_page_break(master)
            composer.append(Document(tmp_path))
        composer.save(output_path)

        for tmp_path in temp_files:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

        msg = f"Success: {total} documents generated"
        if truncated:
            msg += f" (limited to first {MAX_ROWS} of {total_in_file} rows)"
        if not images_by_row:
            msg += " [Note: no photos found in Excel]"
        return True, msg, total

    except Exception as e:
        return False, f"Error: {str(e)}", 0
    finally:
        for p in image_temps:
            try:
                os.remove(p)
            except Exception:
                pass


def get_column_names(excel_path):
    try:
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active
        headers = [str(cell.value).strip() for cell in ws[1] if cell.value is not None]
        row_count = sum(
            1
            for row in ws.iter_rows(min_row=2, values_only=True)
            if not all(c is None or str(c).strip() == "" for c in row)
        )
        wb.close()
        return headers, row_count
    except Exception:
        return [], 0


def get_placeholders(template_path):
    try:
        doc = Document(template_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text += para.text + " "
        found = re.findall(r"\{\{([^}]+)\}\}", text)
        return list(dict.fromkeys(found))
    except Exception:
        return []


# ─── Batch / one-row APIs for Generator UI ───────────────────

def _row_key(headers, row, excel_row_idx):
    """Stable key: first non-empty header value, else Row N."""
    for idx, h in enumerate(headers):
        if idx < len(row) and row[idx] is not None and str(row[idx]).strip() != "":
            return str(row[idx]).strip()
    return "Row %d" % excel_row_idx


def _row_summary(headers, row, limit=6):
    parts = []
    for idx, h in enumerate(headers):
        if idx >= limit:
            break
        val = row[idx] if idx < len(row) else ""
        if val is None:
            val = ""
        val = str(val).strip()
        if not val:
            continue
        if h.strip().upper() in IMAGE_PLACEHOLDERS:
            parts.append(h + "=📷")
        else:
            short = val if len(val) <= 40 else val[:37] + "…"
            parts.append("%s=%s" % (h, short))
    return " · ".join(parts)


def list_data_rows(excel_path):
    """
    JSON list of all data rows for UI skip logic:
    [{ "row": 2, "key": "...", "summary": "..." }, ...]
    """
    import json
    try:
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active
        headers = [str(cell.value).strip() for cell in ws[1] if cell.value is not None]
        out = []
        for excel_row_idx, row in enumerate(
            ws.iter_rows(min_row=2, values_only=True), start=2
        ):
            if all(cell is None or str(cell).strip() == "" for cell in row):
                break
            out.append({
                "row": excel_row_idx,
                "key": _row_key(headers, row, excel_row_idx),
                "summary": _row_summary(headers, row),
            })
        wb.close()
        return json.dumps(out, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)})


def generate_one_row(excel_path, template_path, excel_row_num, output_path):
    """
    Build a single-page docx for one Excel row number (1-based sheet row).
    Returns (ok: bool, message: str)
    """
    image_temps = []
    try:
        excel_row_num = int(excel_row_num)
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active
        headers = [str(cell.value).strip() for cell in ws[1] if cell.value is not None]
        row_values = None
        for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if idx == excel_row_num:
                row_values = row
                break
        wb.close()
        if row_values is None:
            return False, "Row %s not found" % excel_row_num

        mapping = {}
        for i, h in enumerate(headers):
            mapping[h] = row_values[i] if i < len(row_values) else ""

        images_by_row = extract_images_by_row(excel_path)
        image_temps = list(images_by_row.values())
        image_path = images_by_row.get(excel_row_num)

        doc = Document(template_path)
        fill_document(doc, mapping, image_path=image_path)
        doc.save(output_path)
        return True, "OK"
    except Exception as e:
        return False, str(e)
    finally:
        for p in image_temps:
            try:
                os.remove(p)
            except Exception:
                pass


def merge_docx_list(paths_json, output_path):
    """
    Merge a JSON list of docx paths into one multi-page document.
    paths_json: '["/a.docx","/b.docx"]'
    Returns (ok, message)
    """
    import json
    try:
        paths = json.loads(paths_json)
        if not paths:
            return False, "No files to merge"
        existing = [p for p in paths if p and os.path.isfile(p)]
        if not existing:
            return False, "No valid docx paths"
        master = Document(existing[0])
        composer = Composer(master)
        for p in existing[1:]:
            add_page_break(master)
            composer.append(Document(p))
        composer.save(output_path)
        return True, "OK:%d" % len(existing)
    except Exception as e:
        return False, str(e)
